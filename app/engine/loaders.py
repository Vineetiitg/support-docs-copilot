from pathlib import Path

from langchain_core.documents import Document


SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".html", ".htm"}


def load_documents(data_dir: str) -> list[Document]:
    root = Path(data_dir)
    root.mkdir(parents=True, exist_ok=True)

    documents: list[Document] = []
    for path in sorted(root.rglob("*")):
        if not path.is_file() or path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        documents.extend(load_document(path))
    return documents


def load_document(path: Path) -> list[Document]:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return [_text_document(path)]
    if suffix == ".pdf":
        return _pdf_documents(path)
    if suffix == ".docx":
        return [_docx_document(path)]
    if suffix in {".html", ".htm"}:
        return [_html_document(path)]
    return []


def _base_metadata(path: Path) -> dict:
    return {
        "source": path.name,
        "source_path": str(path),
        "file_type": path.suffix.lower().lstrip("."),
    }


def _text_document(path: Path) -> Document:
    return Document(page_content=path.read_text(encoding="utf-8", errors="ignore"), metadata=_base_metadata(path))


def _pdf_documents(path: Path) -> list[Document]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("Install pypdf to ingest PDF files.") from exc

    reader = PdfReader(str(path))
    documents: list[Document] = []
    for index, page in enumerate(reader.pages, start=1):
        metadata = _base_metadata(path)
        metadata["page"] = index
        documents.append(Document(page_content=page.extract_text() or "", metadata=metadata))
    return documents


def _docx_document(path: Path) -> Document:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise RuntimeError("Install python-docx to ingest DOCX files.") from exc

    doc = DocxDocument(str(path))
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
    return Document(page_content=text, metadata=_base_metadata(path))


def _html_document(path: Path) -> Document:
    try:
        from bs4 import BeautifulSoup
    except ImportError as exc:
        raise RuntimeError("Install beautifulsoup4 to ingest HTML files.") from exc

    soup = BeautifulSoup(path.read_text(encoding="utf-8", errors="ignore"), "html.parser")
    for tag in soup(["script", "style"]):
        tag.decompose()
    return Document(page_content=soup.get_text("\n", strip=True), metadata=_base_metadata(path))
